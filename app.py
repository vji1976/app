#!/usr/bin/env python3
"""OFFICE MANAGER v1.0

CODED ON: 
	11.24.2019	
CODED BY: 
	MosesLawn	
CODE ENV: 
	python 3.7 win10x64	
NOTES:
	tkinter.widget.winfo_children()
	tkinter.widget.children.values()
		- to access child widgets of a tkinter widget
		primarily a frame
"""
import sys
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
# 3RD PARTY IMPORTS
from tkcalendar import DateEntry
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
# CUSTOM IMPORTS
import wlib
import data

class App(tk.Tk):
	"""Basic tkinter gui interface inheriting from tk.Tk."""
	def __init__(self):
		super().__init__()
		self.iconbitmap('img/favicon.ico')	# uncomment for titlebar icon
		"""
		self.msgico = Image.open("img/favicon.ico")
		self.msgimg = ImageTk.PhotoImage(self.msgico)		
		self.tk.call('wm', 'iconphoto', self._w, self.msgimg)	# this line changes the icon for 	
																# tk messagebox
		"""
		self.title("Funeral Aid 1000")		
		self.drawmenu()
		self.columnconfigure(0, weight=0)
		self.columnconfigure(1, weight=1)
		self.columnconfigure(2, weight=1)
		self.rowconfigure(1, weight=1)
		self.initvars()
		self.drawgui()
		
	def initvars(self):
		self.dataDict = {}	# this main list will hold all
							# data dictionaries for access later
		# main data dictionary population
		self.srvDict = self.makeDataDict(data.srv_Labels)
		self.dpiDict = self.makeDataDict(data.dpi_Labels)
		self.nokDict = self.makeDataDict(data.nok_Labels)
		self.cemDict = self.makeDataDict(data.cem_Labels)
		self.srvDict["Service Type"].set("Full Funeral")
		# add data dicts
		self.dataDict["Service Information"] = self.srvDict
		self.dataDict["Deceased Personal Information"] = self.dpiDict
		self.dataDict["Next of Kin"] = self.nokDict
		self.dataDict["Cemetery Information"] = self.cemDict
		# status bar text variable for messages
		self.statusTxt = tk.StringVar()
		self.statusTxt.set("Program Running")
		# image paths
		self.fb_img_path = 'img/facebook_xui.png'
		self.tw_img_path = 'img/twitter_xui.png'
		self.pa_img_path = 'img/myparish_xui.png'
		self.ch_img_path = 'img/chrome_xui.png'
		self.yt_img_path = 'img/youtube_xui.png'
		self.fhome_icon = 'img/fhome_icon.png'
		# url paths
		self.fb_url = data.social_paths["fb"]
		self.tw_url = data.social_paths["tw"]
		self.pa_url = data.social_paths["mp"]
		self.ch_url = data.social_paths["ch"]
		self.yt_url = data.social_paths["yt"]	
		
	def drawmenu(self):
		self.menubar = wlib.wMenu(self)
		# file menu commands
		self.menubar.fm.add_command(label="Create Document",
									command=lambda: self.createWordDoc(self.dataDict))
		self.menubar.fm.add_separator()
		self.bind_all('<Control-q>', lambda e: sys.exit())
		self.menubar.fm.add_command(label="Quit", underline=0,
									command=sys.exit, accelerator="Ctrl+Q")		
		# options menu commands
		
		# help menu commands
		self.bind_all('<Control-h>', lambda e: self.m_help_about())
		self.menubar.hm.add_command(label="About", underline=0, 
									command=self.m_help_about, accelerator="Ctrl+H")
		
	def drawgui(self):
		self.style = wlib.wStyle()
		
		''' TOOL BAR FRAME START '''
		''' -------------------- '''
		tcol = ttk.Frame(self)
		tcol.grid(row=0, column=0, columnspan=3, sticky='new', padx=6, pady=2)		
		tbar = ttk.Frame(tcol)
		tbar.grid(row=0, column=0, sticky='nw')		
		
		wlib.drawImgWebLink(tbar, self.fb_img_path, self.fb_url).grid(row=0, column=0, sticky='nw', padx=2)
		wlib.drawImgWebLink(tbar, self.tw_img_path, self.tw_url).grid(row=0, column=1, sticky='nw', padx=2)
		wlib.drawImgWebLink(tbar, self.pa_img_path, self.pa_url).grid(row=0, column=2, sticky='nw', padx=2)
		wlib.drawImgWebLink(tbar, self.ch_img_path, self.ch_url).grid(row=0, column=3, sticky='nw', padx=2)
		wlib.drawImgWebLink(tbar, self.yt_img_path, self.yt_url).grid(row=0, column=4, sticky='nw', padx=2)
				
		''' LEFT COLUMN CONTROLS START '''
		''' -------------------------- '''
		lcol = ttk.Frame(self, width=320, style='column.TFrame')
		lcol.grid(row=1, column=0, sticky='nsew', padx=10, pady=10)
		lcol.grid_propagate(0)				# prevent column from resizing to inner widget size
		lcol.columnconfigure(0, weight=1)	# allows child widgets to fill space of column
		
		# - Funeral Information Frame
		srvFrame = wlib.wLabelFrame(lcol, text="Service Information")
		srvFrame.grid(row=0, column=0, sticky='new')
		# -- Service Number
		srvNumFrame = wlib.drawDictEntry(srvFrame,
										 width=6,
										 keystr="Service Number",
										 datadict=self.srvDict)
		srvNumFrame.grid(row=0, column=0, sticky='nw', padx=4, pady=4)
		# -- Service Type
		srvTypeFrame = wlib.drawRadFrame(srvFrame,
										 wlabel='Service Type',
										 rlabels=data.srv_Types,
										 var=self.srvDict['Service Type'])
		srvTypeFrame.grid(row=1, column=0, sticky='nw', padx=4, pady=4)
		
		# -- Service Date Time Location
		srvDTLFrame = ttk.Frame(srvFrame)
		srvDTLFrame.grid(row=2, column=0, sticky='nw', padx=4, pady=4)
		sdFrame = wlib.drawDate(srvDTLFrame, 
								label='Date',
								var=self.srvDict['Service Date'])
		sdFrame.grid(row=0, column=0, sticky='nw')		
		stFrame = wlib.drawCombo(srvDTLFrame,
								 width=7,
							     label='Time',
							     var=self.srvDict['Service Time'],
							     vals=data.times)
		stFrame.grid(row=0, column=1, sticky='nw', padx=8)		
		slFrame = wlib.drawCombo(srvDTLFrame,
								 width=18,
							     label='Location',
							     var=self.srvDict['Service Location'],
							     vals=data.srv_Places)
		slFrame.grid(row=0, column=2, sticky='nw')
		
		# -- Service Day Celebrant
		srvDCFrame = ttk.Frame(srvFrame)
		srvDCFrame.grid(row=3, column=0, sticky='nw', padx=4, pady=4)
		sdayFrame = wlib.drawCombo(srvDCFrame,
								   width=12,
								   label='Day',
								   var=self.srvDict['Service Day'],
								   vals=data.days)
		sdayFrame.grid(row=0, column=0, sticky='nw')
		scFrame = wlib.drawCombo(srvDCFrame,
								 width=16,
								 label='Celebrant',
								 var=self.srvDict['Celebrant'],
								 vals=data.celebrants)
		scFrame.grid(row=0, column=1, sticky='nw', padx=8)		
		
		# - Funeral Home Contact Phone
		fhcpFrame = ttk.Frame(srvFrame)
		fhcpFrame.grid(row=4, column=0, sticky='w', padx=4, pady=4)
		
		fhFrame = ttk.Frame(fhcpFrame)
		fhFrame.grid(row=0, column=0, columnspan=2, sticky='nw')
		fhImgLabel = wlib.drawImgLabel(fhFrame, self.fhome_icon, size=(32,32))
		fhImgLabel.grid(row=0, column=0, sticky='s', pady=7)
		fhCboFrame = wlib.drawCombo(fhFrame,
								    width=18,
								    label='Funeral Home',
								    var=self.srvDict['Funeral Home'],
								    vals=data.fun_Homes)
		fhCboFrame.grid(row=0, column=1, sticky='nw', padx=8)
		
		fconFrame = ttk.Frame(fhcpFrame)
		fconFrame.grid(row=1, column=0)
		fpFrame = wlib.drawEntry(fconFrame,
								 align='v',
								 width=14,
								 label='Phone', 
								 var=self.srvDict['Contact Phone'])
		fpFrame.grid(row=0, column=0, sticky='nw')
		fcFrame = wlib.drawEntry(fconFrame,
								 align='v',
								 width=26,
							     label='Funeral Contact', 
							     var=self.srvDict['F. H. Contact'])
		fcFrame.grid(row=0, column=1, sticky='nw', padx=8)
				
		
		# - Organist Cantor Servers 
		ocsFrame = ttk.Frame(srvFrame)
		ocsFrame.grid(row=5, column=0, sticky='nw', padx=4, pady=4)
				
		orgFrame = ttk.Frame(ocsFrame)
		orgFrame.grid(row=0, column=0)
		orgEntFrame = wlib.drawEntry(orgFrame,
									 width=24,
								     label='Organist', 
								     align='v', 
								     var=self.srvDict['Organist'])								     
		orgEntFrame.grid(row=0, column=0, sticky='nw', pady=4)
		
		canFrame = ttk.Frame(ocsFrame)
		canFrame.grid(row=0, column=1, sticky='nw', padx=4)
		canEntFrame = wlib.drawEntry(canFrame,
									 width=24,
									 label='Cantor',
									 align='v',
									 var=self.srvDict['Cantor'])
		canEntFrame.grid(row=0, column=0, sticky='nw', pady=4)
		
		svrFrame = ttk.Frame(ocsFrame)
		svrFrame.grid(row=1, column=0, columnspan=2, sticky='nw', pady=4)
		svrRxCtr = 0
		for i in range(0, 3):
			svrLabel = 'Server ' + str(i + 1)
			svrEFrame = wlib.drawEntry(svrFrame,
									   width=30,
									   label=svrLabel,
									   var=self.srvDict[svrLabel])
			svrEFrame.grid(row=svrRxCtr, column=0, pady=4)
			svrRxCtr += 1		
		
		''' MIDDLE COLUMN START '''
		''' ------------------- '''
		mcol = ttk.Frame(self, width=300, style='column.TFrame')
		mcol.grid(row=1, column=1, sticky='nsew', pady=10)
		mcol.grid_propagate(0)				# prevent column from resizing to inner widget size
		mcol.columnconfigure(0, weight=1)	# allows child widgets to fill space of column
		# - Deceased Personal Info Frame
		dpiFrame = wlib.wLabelFrame(mcol, text="Deceased Personal Info")
		dpiFrame.grid(row=0, column=0, sticky='new')
		dpiFrame.columnconfigure(0, weight=1)
		wlib.drawDpiEntries(dpiFrame, self.dpiDict)	
		# - Next of Kin Frame
		nokFrame = wlib.wLabelFrame(mcol, text="Next of Kin")
		nokFrame.grid(row=1, column=0, sticky='new')
		wlib.drawEntries(nokFrame, 	self.nokDict)
				
		''' RIGHT COLUMN START '''
		''' ------------------ '''
		rcol = ttk.Frame(self, width=300, style='column.TFrame')
		rcol.grid(row=1, column=2, sticky='nsew', padx=10, pady=10)
		rcol.grid_propagate(0)				# prevent column from resizing to inner widget size
		rcol.columnconfigure(0, weight=1)	# allows child widgets to fill space of column
		# - Cemetery Frame
		cemFrame = wlib.wLabelFrame(rcol, text="Cemetery ~ Burial Information")
		cemFrame.grid(row=0, column=0, sticky='new')
		cemFrame.columnconfigure(0, weight=1)
		wlib.drawCemEntries(cemFrame, self.cemDict)
		# - Resources Frame
		resFrame = wlib.wLabelFrame(rcol, text="Resources")
		resFrame.grid(row=1, column=0, sticky='new')
		fhweblinkFrame = ttk.Frame(resFrame)
		for link in data.fhome_links:
			weblbl = wlib.drawLinkLabel(fhweblinkFrame, label=link[0], url=link[1])
			weblbl.pack()
		fhweblinkFrame.grid(row=0, column=0, sticky='nw')	
		
		''' STATUS BAR ROW START '''
		''' -------------------- '''
		statusbar = ttk.Label(self, 
							  textvariable=self.statusTxt,
							  style='statusbar.TLabel')
		statusbar.grid(row=2, column=0, columnspan=3, sticky='new')
		
	# - END GUI METHODS - #
	# - FUNCTIONALIT METHODS - #
	def makeDataDict(self, labels):
		dataDict = {k : tk.StringVar() for k in labels}
		return dataDict
		
	def createWordDoc(self, datadict):
		"""
			docx adds an initial paragraph automatically to cells in a table
		"""
		headings = list(self.dataDict.keys())
		data_tables = {}						# will hold tables with a key that is a heading
												# and also a key in data dict
		print(headings)
		
		# -- BEGIN WORD DOCUMENT CODE HERE -- #
		doc = Document()
		
		sections = doc.sections					# only 1 section in our doc
		sections[0].top_margin = Inches(.2)		# sets page top margin
		sections[0].left_margin = Inches(.5)	# increases table width due to autofit
												# using left and right margin
		sections[0].right_margin = Inches(.5)
		
		# HEADING TABLE START #
		head_tbl = doc.add_table(rows=1, cols=2)	# heading data table creation
		head_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
		head_tbl.style = 'Table Grid'
		
		head_cells = head_tbl.rows[0].cells
		head_cells[0].text = "OLOP DEATH REGISTERY WORKSHEET"
		hc1_font = head_cells[0].paragraphs[0].runs[0].font
		hc1_font.name = 'Calibri'
		hc1_font.size = Pt(14)
		hc1_font.bold = True
		hc1_font.underline = True
		hc1_font.color.rgb = RGBColor(0, 0, 255)
		
		# *** SPACER *** #
		doc.add_paragraph()
		# *** SPACER *** #
		
		# BODY TABLE START #
		bod_tbl = doc.add_table(rows=2, cols=2)		# body table for program data		
		bod_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
		# bod_tbl.style = 'Table Grid'				# uncomment to add borders for design		
		
		bod_r1_cells = bod_tbl.rows[0].cells
		bod_r2_cells = bod_tbl.rows[1].cells
		
		# cell 0, 0 (dpi cell)
		dpi_tbl = bod_r1_cells[0].add_table(rows=1, cols=2)
		dpi_tbl.rows[0].cells[0].text = 'Deceased Personal'
		dpi_head_font = dpi_tbl.rows[0].cells[0].paragraphs[0].runs[0].font
		dpi_head_font.name = 'Calibri'
		dpi_head_font.size = Pt(12)
		dpi_head_font.bold = True
		dpi_head_font.color.rgb = RGBColor(0, 0, 0)
		dpi_tbl.style = 'Medium List 1'
		
		# cell 0, 1 (srv cell)
		srv_tbl = bod_r1_cells[1].add_table(rows=1, cols=2)
		srv_tbl.rows[0].cells[0].text = 'Service Information'
		srv_head_font = srv_tbl.rows[0].cells[0].paragraphs[0].runs[0].font
		srv_head_font.name = 'Calibri'
		srv_head_font.size = Pt(12)
		srv_head_font.bold = True
		srv_head_font.color.rgb = RGBColor(0, 0, 0)
		srv_tbl.style = 'Medium List 1'
		
		# cell 1, 0 (nok cell)
		nok_tbl = bod_r2_cells[0].add_table(rows=1, cols=2)
		nok_tbl.rows[0].cells[0].text = 'Next of Kin'
		nok_head_font = nok_tbl.rows[0].cells[0].paragraphs[0].runs[0].font
		nok_head_font.name = 'Calibri'
		nok_head_font.size = Pt(12)
		nok_head_font.bold = True
		nok_head_font.color.rgb = RGBColor(0, 0, 0)
		nok_tbl.style = 'Medium List 1'
		
		# cell 1, 1 (cem cell)
		cem_tbl = bod_r2_cells[1].add_table(rows=1, cols=2)
		cem_tbl.rows[0].cells[0].text = 'Cemetery & Burial Info'
		cem_head_font = cem_tbl.rows[0].cells[0].paragraphs[0].runs[0].font
		cem_head_font.name = 'Calibri'
		cem_head_font.size = Pt(12)
		cem_head_font.bold = True
		cem_head_font.color.rgb = RGBColor(0, 0, 0)
		cem_tbl.style = 'Medium List 1'
		
		data_tables['Deceased Personal Information'] = dpi_tbl
		data_tables['Service Information'] = srv_tbl
		data_tables['Next of Kin'] = nok_tbl
		data_tables['Cemetery Information'] = cem_tbl
		
		
		for lbl, tbl in data_tables.items():
			
			for k, v in self.dataDict[lbl].items():
				row = tbl.add_row()
				row.height = Cm(.75)
				row.cells[0].text = str(k)
				row.cells[1].text = v.get()
				# key font for label cells
				kfont = row.cells[0].paragraphs[0].runs[0].font
				kfont.name = 'Calibri'
				kfont.size = Pt(12)
				kfont.color.rgb = RGBColor(0, 0, 0)
				kfont.bold = True
				# value font for data value cells
				vfont = row.cells[1].paragraphs[0].runs[0].font
				vfont.name = 'Calibri'
				vfont.size = Pt(12)
				vfont.color.rgb = RGBColor(64, 64, 64)
				vfont.bold = True
				
				# row.cells[0].merge(row.cells[1])		
		
		
		bod_r1_cells[0].merge(bod_r2_cells[0])
		bod_r1_cells[1].merge(bod_r2_cells[1])
		
		
		doc.save('test.docx')
		# -- END WORD DOCUMENT CODE HERE -- #
		
		try:
			doc.save("test.docx")
		except PermissionError as pe:
			# this error occurs when document is open in word
			messagebox.showerror("Word Document Error", str(pe))
	
	def openWordDoc(self):
		file_name = filedialog.askopenfilename(initialdir="/",
											   title="Open A Service File",
											   filetypes=(("Word Documents", "*.docx"),
														  ("All Files", "*.*")))
		return file_name
					
		
	def updateStatus(self, msg):
		self.statusTxt.set(msg)
		
	# menu command callback functions
	def m_help_about(self):
		messagebox.showinfo(title="About Money Penny App",
							message="Money Penny Office App v1.0")	
		
if __name__ == '__main__':
	root = App()
	root.geometry("960x720+80+80")
	root.minsize(960,720)
	"""
	root.update_idletasks()				# needed so widgets can be redrawn before next idle
	root.overrideredirect(True)			# removes window decorations (titlebar)
	w = root.winfo_screenwidth()		# get width of main window
	h = root.winfo_screenheight()		# get height of main window
	root.resizable(height=0, width=0)	# user cannot resize window
	"""
	root.mainloop()
