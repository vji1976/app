from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

""" CODE FROM 01/26/2020 
sections[0].top_margin = Cm(1)
		table = doc.add_table(rows=2, cols=2)
		
		dpi_svi_cells = table.rows[0].cells
		dpi_svi_cells[0].text = "Deceased Personal Information"
		dpi_svi_cells[1].text = "Service Information"
		
		nok_cem_cells = table.rows[1].cells
		nok_cem_cells[0].text = "Cemetery Infromation"
		nok_cem_cells[1].text = "Next of Kin"
		
		dpi_tbl = dpi_svi_cells[0].add_table(rows=0, cols=2)
		svi_tbl = dpi_svi_cells[1].add_table(rows=0, cols=2)
		
		for k, v in self.dataDict['Deceased Personal Information'].items():
			row = dpi_tbl.add_row()
			row.cells[0].text = k
			row.cells[1].text = v.get()		
			row.cells[0].paragraphs[0].runs[0].bold = True
			kfont = row.cells[0].paragraphs[0].runs[0].font
			kfont.name = 'Calibri'
			kfont.size = Pt(11)
			vfont = row.cells[1].paragraphs[0].runs[0].font
			vfont.name = 'Calibri'
			vfont.size = Pt(10)
			vfont.color.rgb = RGBColor(0x42, 0x24, 0xE9)
			
		for k, v in self.dataDict['Service Information'].items():
			row = svi_tbl.add_row()
			row.cells[0].text = k
			row.cells[1].text = v.get()		
			row.cells[0].paragraphs[0].runs[0].bold = True
			kfont = row.cells[0].paragraphs[0].runs[0].font
			kfont.name = 'Calibri'
			kfont.size = Pt(11)
			vfont = row.cells[1].paragraphs[0].runs[0].font
			vfont.name = 'Calibri'
			vfont.size = Pt(10)
			vfont.color.rgb = RGBColor(0x42, 0x24, 0xE9)
"""
doc = Document()

# change top margin
sec1 = doc.sections[0]
sec1.top_margin = Inches(.2)

t1 = doc.add_table(rows=1, cols=2)
t1.style = 'Light Grid Accent 1'

t1.cells = t1.rows[0].cells
t1.cells[0].text = "OLOP MORTEM DATA WORKSHEET"
t1_headfont = t1.cells[0].paragraphs[0].runs[0].font
t1_headfont.name = 'Comic Sans MS'
t1_headfont.size = Pt(14)
t1_headfont.color.rgb = RGBColor(0, 255, 0)

t1.cells[1].text = "TABLE 1 ROW 1 COL 2"
t2_headfont = t1.cells[1].paragraphs[0].runs[0].font
t2_headfont.name = 'Segoe UI'
t2_headfont.size = Pt(10)
t2_headfont.color.rgb = RGBColor(0, 0, 255)

t1_row2 = t1.add_row()
row2_cells = t1_row2.cells
row2_cells[0].text = "Service Type:"
in_t1 = row2_cells[1].add_table(rows=1, cols=2)
in_t1.rows[0].cells[0].text = "THIS IS IN A NEW TABLE"

paragraph_format = doc.styles['Normal'].paragraph_format
paragraph_format.space_before = Pt(1)

p1 = doc.add_paragraph()
p1_format = p1.paragraph_format
p1_format.alignment
p1_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.add_run("this is the run in the first paragraph")

r1_font = p1.runs[0].font
r1_font.name = 'Calibri'
r1_font.size = Pt(24)
r1_font.color.rgb = RGBColor(192, 0, 0)

doc.save("tester.docx")
'''
def createWordDoc(self, datadict):
		"""
			docx adds an initial paragraph automatically to cells in a table
		"""
		headings = list(self.dataDict.keys())
		print()
			
		
		doc = Document()
		# doc.add_heading("OLOP FUNERAL SERVICE")
		
		table = doc.add_table(rows=1, cols=2)
		
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = "OLOP FUNERAL SERVICE DATA SHEET"
		hdr_cells[1].text = "Funeral Number: " + self.dataDict[headings[0]]["Service Number"].get()
		
		dpi_svi_cells = table.add_row().cells
		dpi_svi_cells[0].text = headings[0]
		dpi_svi_cells[1].text = headings[1]
		
		nok_cem_cells = table.add_row().cells
		nok_cem_cells[0].text = headings[2]
		nok_cem_cells[1].text = headings[3]
		
		dpi_tbl = dpi_svi_cells[0].add_table(rows=1, cols=2)
		nok_tbl = nok_cem_cells[0].add_table(rows=1, cols=2)
		svi_tbl = dpi_svi_cells[1].add_table(rows=1, cols=2)	
		cem_tbl = nok_cem_cells[1].add_table(rows=1, cols=2)
		
		info_data_tables = [dpi_tbl, svi_tbl, nok_tbl, cem_tbl]
		info_dict = {}
						
		for heading in headings:
			info_dict[heading] = info_data_tables[headings.index(heading)]
		
		from docx.enum.table import WD_ALIGN_VERTICAL
		
		for k in info_dict.keys():
			for label, data in self.dataDict[k].items():
				row = info_dict[k].add_row()
				row.cells[0].text = str(label)
				row.cells[1].text = str(data.get())	
			info_dict[k].rows[0].cells[0].vertical_aligment = WD_ALIGN_VERTICAL.TOP
				
		
		dpi_svi_cells[0].merge(nok_cem_cells[0])
		dpi_svi_cells[1].merge(nok_cem_cells[1])		
				
		"""		
		for k, v in self.dataDict[headings[0]].items():
			row = dpi_tbl.add_row()
			row.cells[0].text = str(k)
			row.cells[1].text = str(v.get())
			row.cells[0].paragraphs[0].runs[0].bold = True
			kfont = row.cells[0].paragraphs[0].runs[0].font
			kfont.name = 'Calibri'
			kfont.size = Pt(11)
			vfont = row.cells[1].paragraphs[0].runs[0].font
			vfont.name = 'Calibri'
			vfont.size = Pt(10)
			vfont.color.rgb = RGBColor(0x42, 0x24, 0xE9)
		"""
			
		doc.save('test.docx')
		
		
		try:
			doc.save("test.docx")
		except PermissionError as pe:
			# this error occurs when document is open in word
			messagebox.showerror("Word Document Error", str(pe))
'''
